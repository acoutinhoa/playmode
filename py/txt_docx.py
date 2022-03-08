import time
start = time.time()

import os
from base import var,cor
from docx import Document

# caminho da pasta do playmode
path='/'.join(os.path.abspath(os.getcwd()).split('/')[:-1])

#-------------------

#https://python-docx.readthedocs.io/en/latest/index.html

# https://www.geeksforgeeks.org/python-working-with-docx-module/

# https://automatetheboringstuff.com/chapter13/

#-------------------

#####################################

# visualizar atributos do objeto
# nao ta funcionando mais :(
def dump(obj):
  for attr in dir(obj):
    print("obj.%s = %r" % (attr, getattr(obj, attr)))

# ver texto sem formatacao
def getText(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

#########################################

# unidades
cm = 72/2.54
mm = cm/10
print('>>> 1px == 1cm')


#########################################

# # # # aberturas = os.path.join(path_txt,'aberturas.txt')
# # # # with open(aberturas, encoding="utf-8") as file:
# # # #     aberturas = file.read()

# # # # # textos
# # # # aberturas=aberturas.split('###')
# # # # aberturas={i:txt.split('___') for i,txt in enumerate(aberturas)}

# fontes_do_pc = ['?',]+installedFonts()

# Variable([
#     dict(name = "tipo", ui = "PopUpButton", args = dict(items = tipos)),
#     dict(name = "altura_randomica", ui = "CheckBox", args = dict(value = True)),
#     dict(name = "autoportante", ui = "CheckBox", args = dict(value = True)),
#     dict(name = "modulor", ui = "CheckBox", args = dict(value = False)),
#     dict(name = "fonte", ui = "PopUpButton", args = dict(items = fontes_do_pc)),
#     dict(name = "cor_suporte", ui="EditText", args=dict(text='k')),
#     dict(name = "alturas", ui = "CheckBox", args = dict(value = False)),
#     dict(name = "gira", ui = "CheckBox", args = dict(value = True)),
#     dict(name = "imagem", ui = "CheckBox", args = dict(value = True)),
#     dict(name = "estampa", ui = "CheckBox", args = dict(value = False)),
# ], globals())
    
# fonte=var(fonte,'HelveticaNeue',lista=fontes_do_pc)
# print('fonte =', fonte)

#########################################

# # print the list of paragraphs in the document
# print('List of paragraph objects:->>>')
# print(doc.paragraphs)
  
# # print the list of the runs 
# # in a specified paragraph
# print('\nList of runs objects in 1st paragraph:->>>')
# print(doc.paragraphs[0].runs)
  
# # print the text in a paragraph 
# print('\nText in the 1st paragraph:->>>')
# print(doc.paragraphs[0].text)
  
# # for printing the complete document
# print('\nThe whole content of the document:->>>\n')
# for para in doc.paragraphs:
#     print(para.text)

#########################################

fontes=[
    'Courier', # 0
    'Courier-Bold', # 1
    'Courier-BoldOblique', # 2
    'Courier-Oblique', # 3
    'CourierNewPS-BoldItalicMT', # 4
    'CourierNewPS-BoldMT', # 5
    'CourierNewPS-ItalicMT', # 6
    'CourierNewPSMT', # 7
]

f_tit={
    'pt':fontes[3],
    'en':fontes[4],
}
f_txt={
    'pt':fontes[7],
    'en':fontes[6],
}
f_txt_it={
    'pt':fontes[6],
    'en':fontes[7],
}
f_txt_bld={
    'pt':fontes[5],
    'en':fontes[4],
}

# fontSize
fs=2
fst=3
#lineHeight
lh=3.5
lht=4
#align
a='left'
at='right'
#paragraphTopSpacing
pts=fs

fmt_tit = FormattedString(
    font=fontes[5],
    fontSize=fst, 
    lineHeight=lht,
    align=at,
)
fmt_txt = FormattedString(
    fontSize=fs, 
    lineHeight=lh,
    align=a,
    paragraphTopSpacing=pts,
)

linguas=['pt','en']

aberturas={}

n=0
T=0
lang=linguas[0]

# docx
doc_path = os.path.join(path,'_/txt/aberturas.docx')
doc = Document(doc_path)

for p,para in enumerate(doc.paragraphs):
    pt=para.text
    if not pt:
        n+=1
        lang=linguas[0]
        T=0

    elif pt=='_':
        lang=linguas[1]
        T=0

    else:
        if n not in aberturas:
            aberturas[n]={}
            for l in linguas:
                aberturas[n][l]={
                    'titulo':fmt_tit.copy(),
                    'texto':fmt_txt.copy(),
                }
        
        if not T:
            tt=aberturas[n][lang]['titulo']
        else:
            tt=aberturas[n][lang]['texto']

        for r,run in enumerate(para.runs):
            t=run.text
            if t:
                if not T:
                    tt.font(f_tit[lang])
                else:
                    if run.bold:
                        tt.font(f_txt_bld[lang])
                        # print('    bold')
                    if run.italic:
                        tt.font(f_txt_it[lang])
                        # print('    italico')
                    else:
                        tt.font(f_txt[lang])
                tt.append(t)
        T+=1
        tt.append('\n')





# for p,para in enumerate(doc.paragraphs):
    
#     # txt curatorial
#     print('##############')
#     print(p)
#     # print(para.paragraph_format.left_indent)

#     # # # for attr in dir(para):
#     # # #     print("obj.%s = %r" % (attr, getattr(para, attr)))

#     for r,run in enumerate(para.runs):

#         txt=run.text
        
#         if txt:
#             print('    ',r)
#             print(txt)
        
#         if run.bold:
#             print('    bold')
#         if run.italic:
#             print('    italico')
#         if run.underline:
#             print('    --------underline')
#         # if str(run.font.color.rgb) == '0070C0':
#         if run.font.color.rgb:
#             print('    >>>>>>>>>>>>>>>>>>>>>>>>>>', run.font.color.rgb)
#         # print(type(run.font.color.rgb))

#         # # # for attr in dir(run):
#         # # #     print("obj.%s = %r" % (attr, getattr(run, attr)))

    
#########################################

end = time.time()
print('\n>>>', end-start, 's')
