import time
start = time.time()

import os
from base import var,cor
from docx import Document

# caminho da pasta do playmode
path='/'.join(os.path.abspath(os.getcwd()).split('/')[:-1])

# unidades
# 1 inch = 2.54 cm
# cm = dpi/2.54
cm = 72/2.54
mm = cm/10

#####################################
# visualizar atributos do objeto
def dump(obj):
  for attr in dir(obj):
    print("obj.%s = %r" % (attr, getattr(obj, attr)))
#####################################

def medidas(painel='',keys=False):
    dados={
        'pg1':{
            'w':460,
            'h':320,
            },
        'pg2':{
            'w':1250,
            'h':320,
            },
        'pg3':{
            'w':600,
            'h':320,
            },
        'painel_playmode':{
            'w':80,
            'h':220,
            'margem':6,
            'dist':4,
            'base':randint(20,60),
            'b2':randint(20,60),
            'suporte':'suspenso',
            },
        'painel_corredor':{
            'w':70,
            'h':220,
            'margem':5,
            'dist':20,
            'base':randint(20,80),
            'b2':randint(20,80),
            'suporte':'suspenso',
            },
        'painel_cortina':{
            'w':80,
            'h':250,
            'margem':0,
            'dist':40,
            'base':randint(20,80),
            'b2':randint(20,80),
            'suporte':'mobile',
            },
        'painel_obra3':{
            'w':30,
            'h':220,
            'margem':2.5,
            'dist':20,
            'base':randint(20,80),
            'b2':randint(20,80),
            'suporte':'mobile',
            },
        'painel_obra2':{
            'w':30,
            'h':220,
            'margem':2.5,
            'dist':50,
            'base':70,
            'b2':None,
            'suporte':'obra2',
            },
        'painel_obra1':{
            'w':30,
            'h':250,
            'margem':2.9,
            'dist':20,
            'base':randint(20,80),
            'b2':randint(20,80),
            'suporte':'suspenso',
            },
        'painel_entrada':{
            'w':20,
            'h':250,
            'margem':0,
            'dist':5,
            'base':5,
            'b2':randint(20,80),
            'suporte':'mobile',
            },
    }
    if painel:
        return dados[painel]
    elif keys:
        return list(dados.keys())

paineis_textos = [
    'painel_playmode',
    'painel_corredor',
]

paineis_obras = [
    'painel_obra1',
    'painel_obra2',
    'painel_obra3',
]


#_________

def tcd_img(img,base,alfa=1):
    imgw,imgh=imageSize(img)
    ei=largura/imgw
    save()
    translate(0,base)
        
    im=ImageObject()
    with im:
        size(imgw,(altura-base)/ei)
        
        if imgh*ei < altura-base:
            fill(1)
            rect(0,0,largura/ei,(altura-base)/ei)
            if painel in paineis_obras:
                    y=(altura-base-67)/ei
            else:
                if pasta[-1] == '0':
                    y=(altura-base-50)/ei
                elif pasta[-1] == '1':
                    y=(altura-base-30)/ei
        else:
            y=0
            
        image(img,(0,y))
    
    scale(ei)
    image(im,(0,0),alpha=alfa)
    restore()

#_________

def giradinha(a):
    eg=1-abs( a/(2*pi) )

    translate(largura/2,0)
    scale(eg,1)
    skew(0, angle2=a, center=(0,0))
    translate(-largura/2,0)
    
    return eg

#_________

def rect_tecido(largura,altura,base,cor_linha=(.6,)):
    save()
    fill(None)
    stroke(*cor_linha)
    strokeWidth(.1)
    rect(0,base,largura,altura-base)
    restore()

#_________

def tecido(largura, altura, base, b2,gira=0,b2c=(0.93,),img=[],alfa=1):
    save()

    if gira:
        a=choice([1,-1]) * randint(0,90)/100 * 2*pi
    else:
        a=0

    giradinha(a)

    #b2
    # if b2 and b2<base:
    if img:
        if alfa <1:
            save()
            translate(largura,0)
            scale(-1,1)
            tcd_img(img[1],b2,alfa=alfa)
            restore()
            rect_tecido(largura,altura,b2)
                        
            fill(1,1,1,alfa)
            rect(0,base,largura,altura-base)
            blendMode('darken')
            
        tcd_img(img[0],base,alfa=1)
        rect_tecido(largura,altura,base)

    else:
        stroke(0.6,)
        strokeWidth(.1)
        if b2:
            fill(*b2c,alfa)
            rect(0,b2,largura,altura-b2)
    
        if painel=='painel_entrada' and l%2:
            fill(.95)
        else:
            fill(1)
        rect(0,base,largura,altura-base)

    restore()
    return a

#_________

def suporte(tipo,largura,altura,ajuste=0,a=1,cor_sup='b'):

    def fio(w,h):
        save()
        eg=giradinha(a)
        w=w/eg
        fill(0)
        rect(largura/2-w/2,altura,w,h)
        restore()
    
    tipo_=tipo.split('_')

    fe=.4 # espessura fio
    be=1 # espessura da barra
    ah=4 # altura do apoio (base/lampada)

    if tipo_[0] == 'base':
        ajuste=10 # +largura do suporte autoportante
    elif tipo == 'suspenso':
        ajuste=5 # +largura do suporte autoportante
    elif tipo == 'obra2':
        ajuste=40 # +largura do suporte autoportante
        ah=5
        
    bw=largura+ajuste # largura do suporte base

    if tipo == 'base_mobile':
        fh=30 # altura do fio para os suportes mobile com base
        bh=altura+fh # altura do suporte base
    else:
        bh=altura
        fh=meio+ph

    sup=BezierPath()
    
    if tipo_[-1] == 'mobile':
        fio(fe,fh)

    if tipo_[0] == 'base':
        # base
        sup.rect(0,0,bw,ah)
        # portico
        sup.rect(0,0,be,bh)
        sup.rect(bw-be,0,be,bh)
        sup.rect(0,bh-be,bw,be)

    elif tipo_[-1] != 'mobile':
        pass
        # portico
        sup.rect(0,bh-be,be,ph)
        sup.rect(bw-be,bh-be,be,ph)
        if tipo == 'obra2':
            sup.rect(0,bh-ah,bw,ah)
        else:
            sup.rect(0,bh-be,bw,be)
    
    # tecido
    if tipo_[-1] != 'mobile':
        if tipo == 'obra2':
            sup.rect(ajuste/2,bh,bw-ajuste,-ah)
        else:
            sup.rect(ajuste/2,bh,bw-ajuste,-be)

    sup.translate(-ajuste/2,0)

    save()
    fill(*cor(cor_sup))
    drawPath(sup)
    restore()

#_________

def formas(caracteres,m,fs=0,fonte='CourierNewPSMT'):
    if not fs:
        fs=m
    base={}
    # formas geometricas
    car_lista=['#','o','t','x',]
    for c in car_lista:
        x,y=0,0
        bezier=BezierPath()
        if c == '#':
            c='quadrado'
            bezier.rect(x,y,m,m)
        elif c == 'o':
            c='circulo'
            bezier.oval(x,y,m,m)
        elif c == 't':
            c='triangulo'
            bezier.polygon((x,y),(x+m/2,y+m),(x+m,y))
        elif c == '+':
            c='cruz'
            bezier.rect(x+m/4,y,m/2,m)
            bezier.rect(x,y+m/4,m,m/2)
        elif c == '*':
            c='xis2'
            n=3
            bezier.polygon((x,y),(x+m/n,y),(x+m,y+m),(x+m-m/n,y+m))
            bezier.polygon((x+m-m/n,y),(x+m,y),(x+m/n,y+m),(x,y+m),)
        elif c == 'X':
            c='xis'
            n=4
            bezier.polygon((x+m/n,y),(x+m,y+m-m/n),(x+m-m/n,y+m),(x,y+m/n))
            bezier.polygon((x,y+m-m/n),(x+m-m/n,y),(x+m,y+m/n),(x+m/n,y+m))
        elif c == 'x':
            c='xis'
            n=6
            x1=BezierPath()
            x2=BezierPath()
            x1.polygon((x+m/n,y),(x+m,y+m-m/n),(x+m-m/n,y+m),(x,y+m/n))
            x2.polygon((x,y+m-m/n),(x+m-m/n,y),(x+m,y+m/n),(x+m/n,y+m))
            bezier=x1.union(x2)
        base[c]=bezier
    # caracteres
    for c in caracteres:
        bezier=BezierPath()
        bezier.textBox(c, (-m/2,-m+fs-m,2*m,2*m), font=fonte, fontSize=fs, align='center',)
        base[c]=bezier
    return base

#_________

def img_faixas(tipo,l,pasta=''):
    path_img = os.path.join(path,'img/expo')
    
    path_img = os.path.join(path_img,pasta)
    img_lista = [img for img in os.listdir(path_img) if img[0] not in ['.','_']]
    img_lista.sort()

    img=choice(img_lista)
    
    if img=='grafico':
        pass
    elif img=='playmode':
        pass
    elif tipo in [4,7]:
        img=os.path.join(path_img,'%s/%s.pdf' % (img,l))
    else:
        img=os.path.join(path_img,img)

    return img

#_________

def limpa_fundo(a=1):
    save()
    giradinha(a)
    fill(*fundo)
    rect(-1,altura,largura+2,ph-altura+meio)
    restore()

#_________
#_________
#_________

# textos

# qctx
def qctx(fs,enter=0,espaco=0):
    x=randint(0,3)
    t=FormattedString()
    if x==0:
        t.append(' ◼︎',fontSize=fs-.2,tracking=.1,font=fontes[0])
    elif x==1:
        t.append('  ● ',fontSize=fs-.3,baselineShift=+.05,tracking=-.1,font=fontes[0])
    elif x==2:
        t.append(' ▴',fontSize=fs+.1,baselineShift=-.1,tracking=-.2,font=fontes[0])
    elif x==3:
        t.append(' ×',fontSize=fs,baselineShift=-.1,font=fontes[0])
    if espaco:
        t.append(' ')
    if enter:
        t.append('\n\n\n\n\n')
    return t

fontes=[
    'Courier', # 0
    'Courier-Bold', # 1
    'Courier-BoldOblique', # 2
    'Courier-Oblique', # 3
    'CourierNewPS-BoldItalicMT', # 4
    'CourierNewPS-BoldMT', # 5
    'CourierNewPS-ItalicMT', # 6
    'CourierNewPSMT', # 7
]

f_tit={
    'pt':fontes[5],
    'en':fontes[4],
}
f_txt={
    'pt':fontes[0],
    'en':fontes[6],
}
f_txt_it={
    'pt':fontes[6],
    'en':fontes[0],
}
f_txt_bld={
    'pt':fontes[1],
    'en':fontes[2],
}

# fontSize
fs=1.8
fst=2.5
#lineHeight
lh=2.5
lht=3
#align
a='left'
at='right'
#paragraphTopSpacing
pts=lh

fmt_tit = FormattedString(
    font=fontes[5],
    fontSize=fst, 
    lineHeight=lht,
    align=at,
)
fmt_txt = FormattedString(
    fontSize=fs, 
    lineHeight=lh,
    align=a,
    paragraphTopSpacing=pts,
)

# # qctx
# qctx=fmt_txt.copy()
# qctx.font(fontes[0])
# qctx.append('◼︎',fontSize=fs-.3,tracking=.1,lineHeight=lh/2)
# qctx.append('●',fontSize=fs-.5,baselineShift=+.05,tracking=-.1,lineHeight=lh/2)
# qctx.append('▴',fontSize=fs+.1,baselineShift=-.1,tracking=-.2,lineHeight=lh/2)
# qctx.append('×',fontSize=fs,baselineShift=-.2,lineHeight=lh/2)
# qctx.append('\n\n\n')

linguas=['pt','en']

txt_aberturas={}

n=0
T=0
lang=linguas[0]

# docx
doc_path = os.path.join(path,'_/txt/aberturas.docx')
doc = Document(doc_path)

for p,para in enumerate(doc.paragraphs):
    pt=para.text
    if not pt:
        n+=1
        lang=linguas[0]
        T=0

    elif pt=='_':
        lang=linguas[1]
        T=0

    else:
        if n not in txt_aberturas:
            txt_aberturas[n]={}
            for l in linguas:
                txt_aberturas[n][l]={
                    'titulo':fmt_tit.copy(),
                    'texto':fmt_txt.copy(),
                }
    
        if not T:
            tt=txt_aberturas[n][lang]['titulo']
        else:
            tt=txt_aberturas[n][lang]['texto']

        for r,run in enumerate(para.runs):
            t=run.text
            if t:
                if not T:
                    tt.font(f_tit[lang])
                    tt.append(t.upper())
                else:
                    if run.bold:
                        tt.font(f_txt_bld[lang])
                        # print('    bold')
                    if run.italic:
                        tt.font(f_txt_it[lang])
                        # print('    italico')
                    else:
                        # tt.baselineShift(0)
                        tt.font(f_txt[lang])
                        tt.fontSize(fs)
                        tt.lineHeight(lh)
                    tt.append(t)
        T+=1
        if doc.paragraphs[p+1].text not in ['','_']:
            tt.append('\n')
        else:
            tt.append(qctx(fs,enter=1,espaco=0))
    
base_obras={
    1:{
        1:{
            'nome':'A taça do mundo é nossa',
            'data':'2018',
            'imagens':['taca.jpeg',],
            'pt':{
                'autoria':[
                    ['Jaime Lauriano','Brasil, 1985'],
                ],
                'info':[
                    'Réplica da taça Jules Rimet fundida em latão e cartuchos de munições utilizadas pela Forças Armadas Brasileiras sobre base de compensado naval',
                    '130 x 30 x 30 cm',
                ],
                'texto':'''A Taça do Mundo é Nossa é uma escultura feita a partir de uma réplica da taça Jules Rimet, usando restos derretidos de cartuchos de munição coletados em áreas de conflito armado no Brasil; especialmente cartuchos usados pelas forças militares. A escolha por fazer uma réplica precisa da taça Jules Rimet parte de dois eixos conceituais principais: o primeiro eixo conceitual foca sua atenção na utilização do futebol como instrumento de propaganda dos regimes militares, que dominaram a América do Sul entre os anos 1960 e 1980, período esse que ficou marcado pela “Operação Condor”, que promovia a interação entre os serviços de inteligência e repressão das ditaduras de Argentina, Brasil, Chile e Uruguai. Por isso, na base da taça encontra-se gravado o nome desses 4 países e os respectivos períodos de duração das ditaduras. O segundo eixo conceitual é baseado na história da própria taça. Marcada por roubos, a taça original desapareceu no Brasil em 1983. Depois de vencer a Copa do Mundo em 1970 e ficar em definitivo com a taça Jules Rimet, o Brasil começou a mostrá- la publicamente. Porém, em 1983 a taça desapareceu e, para surpresa de todos, foi derretida. Anos depois, a FIFA deu ao Brasil uma réplica; então, a taça em exibição na sede da Confederação Brasileira de Futebol (CBF) nada mais é do que uma réplica.'''
            },
            'en':{
                'autoria':[
                    ['Jaime Lauriano','Brazil, 1985'],
                ],
                'info':[
                    'Réplica da taça Jules Rimet fundida em latão e cartuchos de munições utilizadas pela Forças Armadas Brasileiras sobre base de compensado naval',
                    '130 x 30 x 30 cm',
                ],
                'texto':'''A Taça do Mundo é Nossa is a sculpture made from a replica of the Jules Rimet trophy using melted remains of ammo cartridges collected in areas of armed conflict in Brazil, especially the ones used by military forces. The choice to make an accurate copy of the Jules Rimet trophy starts from two main conceptual axes: the first focuses its attention on the use of soccer as propaganda instrument for the military regimes which dominated South America from the 1960s to the1980s, a period marked by the called “Operation Condor”, which promoted the interaction between the intelligence services and the repression of the Argentinian, Brazilian, Chilean and Uruguayan dictatorships. Therefore, the name of these 4 countries is engraved on the base of the trophy, as well as their respective durations. The second conceptual axis is based on the history of the trophy itself. Marked by robberies, the original trophy disappeared in Brazil in 1983. After winning the 1970 World Cup and permanently taking the Jules Rimet trophy, Brazil started to exhibit it publicly. However, in 1983 the trophy disappeared, and, to everyone’s surprise, it was melted. Years later, FIFA gave Brazil a replica. So, the trophy on display at the Brazilian Football Confederation (CBF) headquarters is nothing more than a copy.'''
            },
        },
    },    
}

obras_path = os.path.join(path,'img/obras')

f_tit=fontes[5]
f_tit_data=fontes[7]
f_autor=fontes[4]
f_autor_info=fontes[6]

f_txt={
    'pt':fontes[0],
    'en':fontes[6],
}
f_txt_it={
    'pt':fontes[6],
    'en':fontes[0],
}
f_txt_bld={
    'pt':fontes[1],
    'en':fontes[2],
}

# fontSize
fs=0.9
fst=1.4
#lineHeight
lho=1.3
lht=1.8
#align
a0='right'
a1='left'

fmt_tit = FormattedString(
    font=fontes[5],
    fontSize=fst, 
    lineHeight=lht,
    align=at,
)
fmt_txt = FormattedString(
    fontSize=fs, 
    lineHeight=lho,
    align=a,
)



txt_obras={}

for eixo in base_obras:
    if eixo not in txt_obras:
        txt_obras[eixo]={}

    for n in base_obras[eixo]:
        if n not in txt_obras[eixo]:
            txt_obras[eixo][n]={}
        
        obra=txt_obras[eixo][n]
        b_obra=base_obras[eixo][n]
        
        for l in linguas:
            obra[l]={}
            
            obra[l]['titulo']=fmt_tit.copy()
            obra[l]['titulo'].append(b_obra['nome'].upper(),font=f_tit,align=a0,)
            
            obra[l]['texto']=fmt_txt.copy()

            obra[l]['texto'].append('\n'+b_obra['data']+'\n',font=f_tit_data,align=a0)
            for autor in b_obra[l]['autoria']:
                for i,info in enumerate(autor):
                    if not i:
                        info=info.upper()
                        obra[l]['texto'].append(info,font=f_autor,align=a1)
                        # q=qctx(fs,enter=0,espaco=1)
                        obra[l]['texto'].append(qctx(fs,enter=0,espaco=1))
                    else:
                        obra[l]['texto'].append(info+'\n',font=f_autor_info,align=a1)

            obra[l]['texto'].append('\n\n\n\n')
            for i in b_obra[l]['info']:
                obra[l]['texto'].append(i+'\n',font=f_txt_it[l],align=a0)
            obra[l]['texto'].append('\n\n\n\n'+b_obra[l]['texto'],font=f_txt[l],align=a1)
            obra[l]['texto'].append(qctx(fs,enter=1,espaco=0))

#_________


def desenha_texto():
    hyphenation(False)

    if painel in paineis_textos:
        col=2 # colunas
        ec=2 # entrecolunas

        tit=texto[abertura]['titulo']
        txt=texto[abertura]['texto']
        meio_texto=meio
        line_h=lh

    else:
        col=1 # colunas
        ec=0 # entrecolunas
        
        tit=texto[abertura]['titulo']
        txt=texto[abertura]['texto']
        meio_texto=meio
        line_h=lho
        

    tit.append('\n')

    tw=(largura-2*margem-(col-1)*ec)/col
    tw,th=textSize(txt, width=tw)
    th=(th/col)/2

    twt,tht=textSize(tit, width=tw)

    save()
    translate(margem,meio_texto)

    if estampa:
        me=14
        save()
        fill(1)
        rect(-margem,-th-me,largura,2*th+tht+2*me)
        restore()
    
    for n in range(col):
        txt=textBox(txt,((tw+ec)*n,-th,tw,2*th))

    twt,tht=textSize(tit, width=tw)
    textBox(tit,(0,th,tw,tht))

    # translate(0,th)
    # for n in range(col):
    #     c=line_h
    #     while c < 2*th:
    #         txt=textBox(txt,((tw+ec)*n,-c,tw,line_h))
    #         c+=line_h

    # twt,tht=textSize(tit, width=tw)
    # textBox(tit,(0,0,tw,tht))

    restore()

#_________
#_________

#########################################

# # docx
# doc = os.path.join(path_txt,'Playmode—textos_parede_PT—EN_2022.docx')
# doc = Document(doc)
# dump(doc)

# # textos
# path_txt = os.path.join(path,'_/txt')

# aberturas = os.path.join(path_txt,'aberturas.txt')
# with open(aberturas, encoding="utf-8") as file:
#     aberturas = file.read()
# # textos
# aberturas=aberturas.split('###')
# aberturas={i:txt.split('___') for i,txt in enumerate(aberturas)}

print('>>> 1px == 1cm')

fontes_do_pc = ['?',]+installedFonts()

tipos = [
    '?',
    '0_playmode',
    '1_eixo1',
    '2_eixo2',
    '3_eixo3',
    '4_cortina',
    '5_obras',
    '6_paineis',
    '7_entrada',
    '8_portas',
    ]

Variable([
    dict(name = "txt_escala_real", ui = "CheckBox", args = dict(value = False)),
    dict(name = "tipo", ui = "PopUpButton", args = dict(items = tipos)),
    dict(name = "altura_randomica", ui = "CheckBox", args = dict(value = True)),
    dict(name = "autoportante", ui = "CheckBox", args = dict(value = True)),
    dict(name = "modulor", ui = "CheckBox", args = dict(value = False)),
    dict(name = "fonte", ui = "PopUpButton", args = dict(items = fontes_do_pc)),
    dict(name = "cor_suporte", ui="EditText", args=dict(text='k')),
    dict(name = "alturas", ui = "CheckBox", args = dict(value = False)),
    dict(name = "gira", ui = "CheckBox", args = dict(value = True)),
    dict(name = "imagem", ui = "CheckBox", args = dict(value = True)),
    dict(name = "estampa", ui = "CheckBox", args = dict(value = False)),
    dict(name = "tecido_transparente", ui = "CheckBox", args = dict(value = False)),
], globals())
    
fonte=var(fonte,'HelveticaNeue',lista=fontes_do_pc)
print('fonte =', fonte)

tipo=var(tipo,lista=tipos, tipo='lista')-1

# medidas
meio=140

# pagina:
if tipo in [4,6,8]:
    dados = medidas('pg2')
elif tipo == 7:
    dados = medidas('pg3')
else:
    dados = medidas('pg1')
pw=dados['w']
ph=dados['h']

# painel
if tipo == 0:
    paineis=[('painel_playmode',txt_aberturas[0]),]
elif tipo == 1:
    paineis=[ ('painel_playmode',txt_aberturas[1]),('painel_obra1',txt_obras[1][1]) ]
elif tipo==2:
    paineis=[('painel_corredor',txt_aberturas[2]),('painel_obra2',txt_obras[1][1])]
elif tipo==3:
    paineis=[('painel_corredor',txt_aberturas[3]),('painel_obra3',txt_obras[1][1])]
elif tipo==4:
    paineis=[('painel_cortina',choice([8,])*['',])]
elif tipo==5:
    paineis=[('painel_obra1',txt_obras[1][1]),('painel_obra2',txt_obras[1][1]),('painel_obra3',txt_obras[1][1]),]
elif tipo==6:
    paineis=[
        ('painel_corredor',txt_aberturas[1]),
        ('painel_cortina',['',]), 
        ('painel_obra1',txt_obras[1][1]),
        ('painel_obra2',txt_obras[1][1]),
        ('painel_obra3',txt_obras[1][1]),
        ('painel_obra1',txt_obras[1][1]),
        ('painel_obra2',txt_obras[1][1]),
        ('painel_obra3',txt_obras[1][1]),
        ('painel_entrada',choice([7,])*['',]),
    ]
elif tipo==7:
    paineis=[('painel_entrada',19*['',])]
elif tipo==8:
    paineis=5*[('painel_cortina',['',])]
    print(paineis)

# pasta imagens
if tipo==7:
    pasta='entrada'
elif tipo==4:
    pasta='cortina'
elif tipo==8:
    pasta='portas'
elif tipo==0:
    if estampa:
        pasta='0/estampa'
    else:
        pasta='0/%s' % randint(0,1)


############################

if txt_escala_real:
    e_pg=cm
else:
    e_pg=1
    newPage(pw*e_pg,ph*e_pg)
    scale(e_pg)

    fundo=(0.8,)
    fill(*fundo)
    rect(0,0,pw,ph)

    # meio + alturas
    hs=[meio,]
    if alturas:
        chaves=medidas(keys=True)
        for k in chaves:
            h=medidas(k)['h']
            if h not in hs:
                hs.append(h)

    # alturas
    save()
    fill(None)
    stroke(.4)
    lineDash(1,2)

    for h in hs:
        line((0,h),(width(),h))

    if alturas:
        fs=6
        fill(0,0,1)
        stroke(None)
        fontSize(fs)
        for h in hs:
            text(str(h)+'cm',(fs/2,h-1.5*fs))        
    restore()

for i,info in enumerate(paineis):
    painel,texto=info
    
    dados = medidas(painel)
    ajuste=0
    
    largura=dados['w']
    altura=dados['h']
    margem=dados['margem']
    dist=dados['dist']
    base=dados['base']
    b2=dados['b2']
    suporte_tipo=dados['suporte']
    

    #__________________________

    # só tecido
    if txt_escala_real:

        if tipo==7:
            autoportante=False
            camadas=4
        else:
            camadas=1
    
        for camada in range(camadas):

            for l,abertura in enumerate(texto):
                if tipo==7 and not (camada+l)%2:
                    desenha=0
                else:
                    desenha=1

                if desenha:
                    if altura_randomica:
                        base=medidas(painel)['base']
                        b2=medidas(painel)['b2']
        
                    newPage(largura*e_pg,(altura-base)*e_pg)
                    scale(e_pg)
                    translate(0,-base)

                    #_________

                    #tecido
                    if tipo==7 and not (camada+l)%2:
                        desenha=0
                    else:
                        desenha=1

                    if desenha:
                        img=[]
                        if imagem and tipo in [4,7,0,8]:
                            img.append(img_faixas(tipo,l,pasta=pasta))
                        print(i,camada,l,img)
                        a=tecido(largura, altura, base, b2, gira=0,img=img,alfa=1)

                    #_________

                    #texto
                    if texto:
                        desenha_texto()
                    #_________
    #__________________________

    else:
        save()
    
        if tipo==5:
            translate(100*(i+1),0)
        elif tipo==8:
            translate(200*(i+1),0)
        elif tipo==6:
            x0=90
            if not i:
                translate(x0,0)
            elif i==1:
                translate(80*(i+1)+x0*2,0)
            elif i in [5,6,7]:
                autoportante=False
                translate(70*(i+1)+x0*3.5,0)
            elif i==8:
                autoportante=False
                translate(70*(i+1)+x0*4,0)
            else:
                translate(70*(i+1)+x0*3.1,0)
            if alturas:
                save()
                translate(0,meio)
                rotate(90)
                fontSize(6)
                fill(1,0,0)
                text(paineis[i][0],(2,10))
                restore()
        
        elif painel in paineis_textos:
            translate(50,0)
        elif painel in paineis_obras:
            translate(310,0)
        else: 
            #centraliza
            tudo=len(texto)*(largura+dist)-dist
            translate((pw-tudo)/2,0)
        
        if tipo==7:
            autoportante=False
            camadas=4
        else:
            camadas=1
    
        if autoportante:
            suporte_tipo='base_'+suporte_tipo
            ajuste=10
            dist+=ajuste

        #=======================
    
        total=len(texto)
        total=total*largura + (total-1)*dist
    
        p=.98 # escala da perspectiva
    
        for camada in range(camadas):
        
            ep = p**(camadas-camada-1)
        
            save()

            #perspectiva
            translate(total/2,meio)
            scale(ep)        
            translate(-total/2,-meio)

            for l,abertura in enumerate(texto):
                if altura_randomica:
                    base=medidas(painel)['base']
                    b2=medidas(painel)['b2']
        
                if gira and suporte_tipo.split('_')[-1]=='mobile':
                    g=1
                else:
                    g=0
        
                #_________

                #tecido
                if tipo in [4,] or tecido_transparente: # tecido transparente
                    n=2
                    alfa=.5
                else:
                    n=1
                    alfa=1
            
                if tipo==7 and not (camada+l)%2:
                    desenha=0
                elif tipo in [5,6] and painel in paineis_obras and l>0:
                    desenha=0
                else:
                    desenha=1

                if desenha:
                    if painel in paineis_obras:
                        pasta='obras/1/1' 
                    img=[]
                    if imagem and tipo in [4,7,0,8,5]:
                        for n in range(n):
                            img.append(img_faixas(tipo,l,pasta=pasta))
                    print(i,camada,l,img)
                    a=tecido(largura, altura, base, b2, gira=g,img=img,alfa=alfa)

                    #_________

                    # suporte
                    suporte(suporte_tipo,largura,altura,ajuste,a=a,cor_sup=cor_suporte)

                    #_________

                    #texto
                    if painel in paineis_textos+paineis_obras:
                        print('>>>>>>',texto)
                        desenha_texto()
            
                #_________
            
                translate(largura+dist,0)

            restore()

        #=======================

        restore()

# modulor
if modulor:
    blendMode('darken')
    for i in range(randint(1,1)):
        modulor = os.path.join(path,'img/escala/modulor%s.pdf' % randint(0,3))
        save()
        translate(randint(220,pw-120),0)
        scale(choice([-1,1]),1)
        image(modulor,(0,0))
        restore()


#########################################

end = time.time()
print('\n>>>', end-start, 's')
